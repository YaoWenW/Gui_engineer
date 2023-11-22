# -*- coding: utf-8 -*-
import sys

from PyQt5.QtGui import QIcon

import Main_window as main_win
import qsm_n
import sj_n
import table
import sj_w
import qsm_w
import win_sub
import win_obj
import sj_win_sub
import sj_win_obj
import demo_main as demo
import numpy as np

from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog
from PyQt5 import QtCore
import qdarkstyle
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class main_win_Ui(QMainWindow, main_win.Ui_MainWindow):
    def __init__(self):
        super(main_win_Ui, self).__init__()
        self.setupUi(self)


class sj_w_Ui(QMainWindow, sj_w.Ui_Dialog):
    def __init__(self):
        super(sj_w_Ui, self).__init__()
        self.setupUi(self)


class qsm_w_Ui(QMainWindow, qsm_w.Ui_Dialog):
    def __init__(self):
        super(qsm_w_Ui, self).__init__()
        self.setupUi(self)


class table_Ui(QMainWindow, table.Ui_Dialog):
    def __init__(self):
        super(table_Ui, self).__init__()
        self.setupUi(self)


if __name__ == '__main__':
    QtCore.QCoreApplication.setAttribute(QtCore.Qt.AA_EnableHighDpiScaling, True)

    app = QApplication(sys.argv)
    app.setStyleSheet(qdarkstyle.load_stylesheet_pyqt5())
    main_win = main_win_Ui()
    main_win.setWindowIcon(QIcon('app_icon.ico'))
    main_win.show()

    sj_w = sj_w_Ui()
    qsm_w = qsm_w_Ui()
    win_sub = win_sub.MatrixInputApp()
    # win_obj = win_obj_old.DynamicGridLayoutWidget()
    win_obj = win_obj.MatrixCalculator()
    sj_win_sub = sj_win_sub.MatrixInputApp()

    # sj_win_obj = sj_win_obj_old.DynamicGridLayoutWidget()
    sj_win_obj = sj_win_obj.MatrixCalculator()
    qsm_n = qsm_n.MatrixCalculator()
    sj_n = sj_n.MatrixCalculator()
    table = table_Ui()


    def sub_ww():
        try:
            x = qsm_w
            zg = [x.zg11, x.zg12, x.zg13, x.zg14, x.zg15, x.zg16, x.zg17,
                  x.zg21, x.zg22, x.zg23,
                  x.zg31, x.zg32, x.zg33, x.zg34,
                  x.zg41, x.zg42, x.zg43, x.zg44,
                  x.zg51, x.zg52]
            sub = win_sub.sub_w[5:]
            for i in range(len(sub)):
                zg[i].setText(str(sub[i]))
            win_sub.close()
        except:
            print('error')
        return


    def obj_ww():
        try:
            x = qsm_w
            kg = [x.kg11, x.kg12, x.kg13, x.kg14, x.kg15, x.kg16, x.kg17,
                  x.kg21, x.kg22, x.kg23,
                  x.kg31, x.kg32, x.kg33, x.kg34,
                  x.kg41, x.kg42, x.kg43, x.kg44,
                  x.kg51, x.kg52]
            obj = win_obj.obj_w
            for i in range(len(obj)):
                kg[i].setText(str(obj[i]))
            win_obj.close()
        except:
            print('error')
        return


    def zh_ww():
        try:
            x = qsm_w
            zh = [x.zh11, x.zh12, x.zh13, x.zh14, x.zh15, x.zh16, x.zh17,
                  x.zh21, x.zh22, x.zh23,
                  x.zh31, x.zh32, x.zh33, x.zh34,
                  x.zh41, x.zh42, x.zh43, x.zh44,
                  x.zh51, x.zh52]
            zhw = demo.combine2_W(win_sub.sub_w, win_obj.obj_w)
            for i in range(len(zhw)):
                zh[i].setText(str(zhw[i]))
        except:
            print('error')
        return


    def save_ww():
        qsm_w.close()
        main_win.label_qsm_w.setText("组合权重计算成功！")


    def save_nn():
        qsm_n.close()
        main_win.label_qsm_n.setText("确定度计算成功！")


    def res():
        try:
            zhw = demo.combine2_W(win_sub.sub_w, win_obj.obj_w)
            L, level = demo.cal_L(qsm_n.matrix, zhw)
            main_win.lable_qsm_level.setText(f"隶属度:{L},\n可靠性:{level}")
        except:
            main_win.lable_qsm_level.setText("请先完成上面的步骤！")
        return


    def save():
        try:
            doc = Document()
            project = main_win.lineEdit_project.text()
            name = main_win.lineEdit_name.text()  # 负责人
            start = main_win.lineEdit.text()  # 开工时间
            end = main_win.lineEdit_2.text()  # 竣工时间
            run = main_win.lineEdit_3.text()  # 开始运营时间

            sub_input = win_sub.matrices  # 比较矩阵
            sub = win_sub.sub_w  # 主观权重
            X = win_obj.input_values  # 客观权重打分值
            obj_input = [[round(X[j] / X[i], 3) for i in range(len(X))] for j in range(len(X))]
            obj = win_obj.obj_w  # 客观权重
            zhw = demo.combine2_W(sub, obj)  # 组合权重
            n_input = qsm_n.input_values
            n_matrix = qsm_n.matrix  # 单因素矩阵
            L, level = demo.cal_L(qsm_n.matrix, zhw)
            text = '\n'.join(["全寿命阶段", f"项目名称：{project}", f"负责人：{name}",
                              f"开工时间：{start}，竣工时间：{end}，开始运营时间：{run}",
                              f"比较矩阵：{sub_input}", f"主观权重：{sub}",
                              f"客观权重计算输入：{obj_input}", f"客观权重：{obj}", f"组合权重：{zhw}",
                              f"确定度计算输入：{n_input}", f"单因素矩阵：{n_matrix}",
                              f"隶属度：{L}", f"可靠性等级：{level}"])
            print(text)
            doc.add_paragraph().add_run("全寿命阶段").bold = True
            text = '\n'.join([f"项目名称：{project}", f"负责人：{name}",
                              f"开工时间：{start}，竣工时间：{end}，开始运营时间：{run}",
                              f"可靠性等级：{level}"])
            doc.add_paragraph(text)

            # 定义表格数据

            doc.add_paragraph().add_run("步骤1：权重计算").bold = True
            doc.add_paragraph().add_run("1. 主观权重计算：\n        将问卷调查的数据进行整理作为主观权重计算的输入数据，得到比较矩阵如表1-表6所示")

            # 比较矩阵
            zb_name = ['材料因素U1', '结构抗变形能力因素U2', '水文地质条件因素U3',
                       '施工期影响因素','运营期影响因素',
                       '管片材料U11', '管片质量U12', '防水材料类型U13', '防水材料老化特性U14', '密封垫构型/防水能力U15', '密封垫固定形式U16', '密封垫布置形式U17',
                       '纵缝接头刚度U21', '环缝抗剪刚度U22', '内部结构类型U23',
                       '水压U31', '不均匀地层U32', '温度影响U33', '水位影响U34',
                       '施工期上浮变形U41', '拼装精度U42', '姿态控制U43', '成型管片U44',
                       '水位影响U51', '纵向应力松弛U52'
                       ]
            hzb1 = zzb1 = zb_name[:5]
            hzb2 = zzb2 = zb_name[5:12]
            hzb3 = zzb3 = zb_name[12:15]
            hzb4 = zzb4 = zb_name[15:19]
            hzb5 = zzb5 = zb_name[19:23]
            hzb6 = zzb6 = zb_name[23:]

            hzb_1 = [hzb1, hzb2, hzb3, hzb4,hzb5,hzb6]
            zzb_1 = [zzb1, zzb2, zzb3, zzb4,zzb5,zzb6]
            data_1 = sub_input

            table_name = ['表1 准则层评价指标比较矩阵A', '表2 材料因素U1评价指标比较矩阵A', '表3 结构抗变形能力因素U2评价指标比较矩阵A',
                          '表4 水文地质条件因素U3评价指标比较矩阵A', '表5 施工期影响因素U4评价指标比较矩阵A', '表6 运营期影响因素U5评价指标比较矩阵A']
            for i in range(len(table_name)):
                create_table_with_title(doc, table_name[i], hzb_1[i], zzb_1[i], len(data_1[i]), len(data_1[i][0]),
                                        data_1[i])

            # 主观权重
            doc.add_paragraph().add_run('通过改进的层次分析法可以得到盾构隧道全寿命阶段防水可靠性评价指标的主观权重值如表7所示。')
            # sub = [9.876, 8.765, 7.654,6.543, 5.432, 4.321, 3.210, 2.109, 1.098]
            data_2 = [[x] for x in sub[5:]]
            hzb_2 = ['影响因素', '主观权重']
            zzb_2 = zb_name[5:]

            create_table_with_title(doc, '表7 盾构隧道全寿命阶段防水可靠性评价指标主观权重值', hzb_2, zzb_2, len(data_2), len(data_2[0]), data_2)
            # 客观权重
            doc.add_paragraph().add_run('2. 客观权重计算：\n        采用层次分析法原理将每个评价指标因素对防水能力的影响程度进行两两对比得到评判矩阵如表8所示')
            # obj_input = [[1, 1], [1, 1]]
            data_3 = obj_input
            hzb_3 = zzb_3 = zb_name[5:]
            zzb_3 = [i[-3:] for i in zzb_3]
            create_table_with_title(doc, '表8 评价指标因素对防水能力影响程度的评判矩阵', hzb_3, zzb_3, len(data_3), len(data_3[0]), data_3)

            doc.add_paragraph().add_run(
                '采用求解矩阵最大特征根的方法计算判断矩阵的最大特征值及其对应的特征向量，并对最大特征向量进行归一化处理得到盾构隧道全寿命阶段防水可靠性评价指标客观权重值如表9所示')
            # obj = [9.876, 8.765, 7.654, 6.543, 5.432, 4.321, 3.210, 2.109, 1.098]
            data_4 = [[x] for x in obj]
            hzb_4 = ['影响因素', '客观权重']
            zzb_4 = zb_name[5:]
            create_table_with_title(doc, '表9 盾构隧道全寿命阶段防水可靠性评价指标客观权重值', hzb_4, zzb_4, len(data_4), len(data_4[0]), data_4)

            # 组合权重
            doc.add_paragraph().add_run(
                f'3. 组合权重计算\n        根据博弈论法对得到的主、客观权重进行综合计算得到组合权重。以将组合权重与主、客观权重之间的离差极小化作为目标，计算得到的主观权重如表10所示。')
            # zh_w = [9.876, 8.765, 7.654, 6.543, 5.432, 4.321, 3.210, 2.109, 1.098]
            data_5 = np.array([sub[5:], obj, zhw]).T
            hzb_5 = ['主观权重', '客观权重', '组合权重']
            zzb_5 = zb_name[5:]
            create_table_with_title(doc, '表10 盾构隧道全寿命阶段防水可靠性评价指标主、客观权重值', hzb_5, zzb_5, len(data_5), len(data_5[0]),
                                    data_5)
            # 单因素确定度计算
            doc.add_paragraph().add_run("步骤二：单因素确定度计算").bold = True
            doc.add_paragraph().add_run("根据需要评价项目的实际情况得到盾构隧道防水全寿命阶段可靠性评价指标的实际取值如表11所示")
            # n_input = [9.876, 8.765, 7.654, 6.543, 5.432, 4.321, 3.210, 2.109, 1.098]
            data_6 = [[x] for x in n_input]
            hzb_6 = ['影响因素', '取值']
            zzb_6 = zb_name[5:]
            create_table_with_title(doc, '表11 盾构隧道全寿命阶段防水可靠性评价指标实际取值', hzb_6, zzb_6, len(data_6), len(data_6[0]), data_6)

            # n_matrix = [[1, 2, 3], [4, 5, 6], [7, 8, 9]]
            data_7 = n_matrix
            hzb_7 = ['一级', '二级', '三级', '四级', '五级']
            zzb_7 = zb_name[5:]
            zzb_7 = [i[-3:] for i in zzb_7]
            create_table_with_title(doc, '表12 盾构隧道全寿命阶段防水可靠性评价指标单因素确定度', hzb_7, zzb_7, len(data_7), len(data_7[0]),
                                    data_7)

            # 隶属度
            doc.add_paragraph().add_run("步骤三：盾构隧道全寿命阶段防水体系可靠性等级计算").bold = True
            doc.add_paragraph().add_run("根据计算得到的组合权重和单因素确定度，采用加权平均法通过公式将地层评价结果传递给上一层，进而得到盾构隧道设计阶段防水可靠性评价等级结果如表13所示")
            # L = [0.1,0.1,0.1,0.1,0.1]
            data_8 = [L]
            hzb_8 = ['一级（低可靠性）', '二级（较低可靠性）', '三级（中等可靠性）', '四级（较高可靠性）', '五级（高可靠性）']
            zzb_8 = ['隶属度计算值']
            create_table_with_title(doc, '表13 盾构隧道全寿命阶段防水可靠性评价等级结果', hzb_8, zzb_8, len(data_8), len(data_8[0]),
                                    data_8)

            doc.add_paragraph().add_run(f'根据计算得到的综合确定度，按照隶属度最大原则得到最终的隶属度等级为{level}')

            options = QFileDialog.Options()
            file_name, _ = QFileDialog.getSaveFileName(caption='Save File', directory='',
                                                       filter="Word Files (*.docx);;All Files (*)", options=options)

            doc.save(file_name)
            main_win.lable_qsm_save.setText("导出成功！")
        except:
            main_win.lable_qsm_save.setText("导出失败，请按照要求完成操作！")
        return


    def sj_sub_ww():
        try:
            x = sj_w
            zg = [x.zg11, x.zg12, x.zg13, x.zg14, x.zg15, x.zg16, x.zg17,
                  x.zg21, x.zg22, x.zg23,
                  x.zg31, x.zg32, x.zg33, x.zg34]
            sub = sj_win_sub.sub_w[3:]
            for i in range(len(sub)):
                zg[i].setText(str(sub[i]))
            sj_win_sub.close()
        except:
            print("error")
        return


    def sj_obj_ww():
        try:
            x = sj_w
            kg = [x.kg11, x.kg12, x.kg13, x.kg14, x.kg15, x.kg16, x.kg17,
                  x.kg21, x.kg22, x.kg23,
                  x.kg31, x.kg32, x.kg33, x.kg34]
            obj = sj_win_obj.obj_w
            for i in range(len(obj)):
                kg[i].setText(str(obj[i]))
            sj_win_obj.close()
        except:
            print("error")
        return


    def sj_zh_ww():
        try:
            x = sj_w
            zh = [x.zh11, x.zh12, x.zh13, x.zh14, x.zh15, x.zh16, x.zh17,
                  x.zh21, x.zh22, x.zh23,
                  x.zh31, x.zh32, x.zh33, x.zh34]
            zhw = demo.combine2_W(sj_win_sub.sub_w, sj_win_obj.obj_w)
            for i in range(len(zhw)):
                zh[i].setText(str(zhw[i]))
        except:
            print('error')
        return


    def sj_save_ww():
        sj_w.close()
        main_win.label_5.setText("权重计算成功！")
        return

    def sj_save_nn():
        sj_n.close()
        main_win.label_6.setText("确定度计算成功！")
        return


    def sj_res():
        try:
            zhw = demo.combine2_W(sj_win_sub.sub_w, sj_win_obj.obj_w)
            L, level = demo.cal_L(sj_n.matrix, zhw)
            main_win.lable_sj_level.setText(f"隶属度:{L},\n可靠性:{level}")
        except:
            main_win.lable_sj_level.setText("请先完成上面步骤！")
        return


    def create_table_with_title(doc, name, hzb, zzb, rows, cols, data):
        # 表名
        paragraph = doc.add_paragraph()
        paragraph.add_run(name).bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_after = Pt(0)

        # 创建表格
        table = doc.add_table(rows=rows + 1, cols=cols + 1)
        table.style = 'Table Grid'

        # 设置表格中的横坐标
        for i in range(cols):
            cell = table.cell(0, i + 1)
            cell.text = hzb[i]
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # for i in range(rows+1):
        #     cell = table.cell(i, 0)
        #     cell.width = Cm(5)
        # 设置表格中的纵坐标
        for i in range(rows):
            cell = table.cell(i + 1, 0)
            cell.text = zzb[i]

        # 设置表格内容
        for i in range(rows):
            for j in range(cols):
                cell = table.cell(i + 1, j + 1)
                if i < len(data) and j < len(data[i]):
                    cell.text = str(data[i][j])
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        # "{:.3f}".format(data[i][j])

        doc.add_paragraph("")


    def sj_save():
        try:
            doc = Document()
            project = main_win.lineEdit_project.text()
            name = main_win.lineEdit_name.text()  # 负责人
            start = main_win.lineEdit.text()  # 开工时间
            end = main_win.lineEdit_2.text()  # 竣工时间
            run = main_win.lineEdit_3.text()  # 开始运营时间

            sub_input = sj_win_sub.matrices  # 比较矩阵
            sub = sj_win_sub.sub_w # 主观权重
            X = sj_win_obj.input_values  # 客观权重打分值
            obj_input = [[round(X[j] / X[i], 3) for i in range(len(X))] for j in range(len(X))]
            obj = sj_win_obj.obj_w  # 客观权重
            zhw = demo.combine2_W(sub, obj)  # 组合权重
            n_input = sj_n.input_values
            n_matrix = sj_n.matrix  # 单因素矩阵
            L, level = demo.cal_L(sj_n.matrix, zhw)

            text = '\n'.join([f"项目名称：{project}", f"负责人：{name}",
                              f"开工时间：{start}，竣工时间：{end}，开始运营时间：{run}",
                              f"比较矩阵：{sub_input}", f"主观权重：{sub}",
                              f"客观权重计算输入：{obj_input}", f"客观权重：{obj}", f"组合权重：{zhw}",
                              f"确定度计算输入：{n_input}", f"单因素矩阵：{n_matrix}",
                              f"隶属度：{L}", f"可靠性等级：{level}"])
            print(text)

            doc.add_paragraph().add_run("设计阶段").bold = True
            # level = '5级'
            text = '\n'.join([f"项目名称：{project}", f"负责人：{name}",
                              f"开工时间：{start}，竣工时间：{end}，开始运营时间：{run}",
                              f"可靠性等级：{level}"])
            doc.add_paragraph(text)

            # 定义表格数据


            doc.add_paragraph().add_run("步骤1：权重计算").bold = True
            doc.add_paragraph().add_run("1. 主观权重计算：\n        将问卷调查的数据进行整理作为主观权重计算的输入数据，得到比较矩阵如表1-表4所示")

            # 比较矩阵
            zb_name = ['材料因素U1', '结构抗变形能力因素U2', '水文地质条件因素U3',
                       '管片材料U11', '管片质量U12', '防水材料类型U13', '防水材料老化特性U14', '密封垫构型/防水能力U15', '密封垫固定形式U16', '密封垫布置形式U17',
                       '纵缝接头刚度U21', '环缝抗剪刚度U22', '内部结构类型U23',
                       '水压U31', '不均匀地层U32', '温度影响U33', '水位影响U34',
                       ]
            hzb1 = zzb1 = zb_name[:3]
            hzb2 = zzb2 = zb_name[3:10]
            hzb3 = zzb3 = zb_name[10:13]
            hzb4 = zzb4 = zb_name[13:]

            hzb_1 = [hzb1, hzb2, hzb3, hzb4]
            zzb_1 = [zzb1, zzb2, zzb3, zzb4]
            # sub_input = [[[1, 1], [1, 1]], [[1, 1], [1, 1]], [[1, 1], [1, 1]], [[1, 1], [1, 1]]]
            # data1, data2, data3, data4 = sub_input
            data_1 = sub_input

            table_name = ['表1 准则层评价指标比较矩阵A', '表2 材料因素U1评价指标比较矩阵A', '表3 结构抗变形能力因素U2评价指标比较矩阵A', '表4 水文地质条件因素U3评价指标比较矩阵A']
            for i in range(len(table_name)):
                create_table_with_title(doc, table_name[i], hzb_1[i], zzb_1[i], len(data_1[i]), len(data_1[i][0]),
                                        data_1[i])

            # 主观权重
            doc.add_paragraph().add_run('通过改进的层次分析法可以得到盾构隧道设计阶段防水可靠性评价指标的主观权重值如表5所示。')
            # sub = [9.876, 8.765, 7.654,6.543, 5.432, 4.321, 3.210, 2.109, 1.098]
            data_2 = [[x] for x in sub[3:]]
            hzb_2 = ['影响因素','主观权重']
            zzb_2 = zb_name[3:]

            create_table_with_title(doc, '表5 盾构隧道设计阶段防水可靠性评价指标主观权重值', hzb_2, zzb_2, len(data_2), len(data_2[0]), data_2)
            # 客观权重
            doc.add_paragraph().add_run('2. 客观权重计算：\n        采用层次分析法原理将每个评价指标因素对防水能力的影响程度进行两两对比得到评判矩阵如表6所示')
            # obj_input = [[1, 1], [1, 1]]
            data_3 = obj_input
            hzb_3 = zzb_3 = zb_name[3:]
            zzb_3 = [i[-3:] for i in zzb_3]
            create_table_with_title(doc, '表6 评价指标因素对防水能力影响程度的评判矩阵', hzb_3, zzb_3, len(data_3), len(data_3[0]), data_3)

            doc.add_paragraph().add_run('采用求解矩阵最大特征根的方法计算判断矩阵的最大特征值及其对应的特征向量，并对最大特征向量进行归一化处理得到盾构隧道设计阶段防水可靠性评价指标客观权重值如表7所示')
            # obj = [9.876, 8.765, 7.654, 6.543, 5.432, 4.321, 3.210, 2.109, 1.098]
            data_4 = [[x] for x in obj]
            hzb_4 = ['影响因素', '客观权重']
            zzb_4 = zb_name[3:]
            create_table_with_title(doc, '表7 盾构隧道设计阶段防水可靠性评价指标客观权重值', hzb_4, zzb_4, len(data_4), len(data_4[0]), data_4)

            # 组合权重
            # n1 = '公式1'
            doc.add_paragraph().add_run(f'3. 组合权重计算\n        根据博弈论法对得到的主、客观权重进行综合计算得到组合权重。以将组合权重与主、客观权重之间的离差极小化作为目标，计算得到的组合权重如表8所示。')
            # zh_w = [9.876, 8.765, 7.654, 6.543, 5.432, 4.321, 3.210, 2.109, 1.098]
            data_5 = np.array([sub[3:], obj, zhw]).T
            hzb_5 = ['主观权重','客观权重', '组合权重']
            zzb_5 = zb_name[3:]
            create_table_with_title(doc, '表8 盾构隧道设计阶段防水可靠性评价指标主、客观权重值', hzb_5, zzb_5, len(data_5), len(data_5[0]), data_5)

            # 单因素确定度计算
            doc.add_paragraph().add_run("步骤二：单因素确定度计算").bold = True
            doc.add_paragraph().add_run("根据需要评价项目的实际情况得到盾构隧道防水设计阶段可靠性评价指标的实际取值如表9所示")
            # n_input = [9.876, 8.765, 7.654, 6.543, 5.432, 4.321, 3.210, 2.109, 1.098]
            data_6 = [[x] for x in n_input]
            hzb_6 = ['影响因素', '取值']
            zzb_6 = zb_name[3:]
            create_table_with_title(doc, '表9 盾构隧道设计阶段防水可靠性评价指标实际取值', hzb_6, zzb_6, len(data_6), len(data_6[0]), data_6)


            # n_matrix = [[1, 2, 3], [4, 5, 6], [7, 8, 9]]
            data_7 = n_matrix
            doc.add_paragraph().add_run("通过公式求解每个评价指标隶属于某一层次的单因素确定性程度，计算得到的结果如表10所示")
            hzb_7 = ['一级','二级','三级','四级','五级']
            zzb_7 = zb_name[3:]
            zzb_7 = [i[-3:] for i in zzb_7]
            create_table_with_title(doc, '表10 盾构隧道设计阶段防水可靠性评价指标单因素确定度', hzb_7, zzb_7, len(data_7), len(data_7[0]),data_7)

            # 隶属度
            doc.add_paragraph().add_run("步骤三：盾构隧道设计阶段防水体系可靠性等级计算").bold = True
            doc.add_paragraph().add_run("根据计算得到的组合权重和单因素确定度，采用加权平均法通过公式将地层评价结果传递给上一层，进而得到盾构隧道设计阶段防水可靠性评价等级结果如表11所示")
            data_8 = [L]
            hzb_8 = ['一级（低可靠性）', '二级（较低可靠性）', '三级（中等可靠性）', '四级（较高可靠性）', '五级（高可靠性）']
            zzb_8 = ['隶属度计算值']
            create_table_with_title(doc, '表11 盾构隧道设计阶段防水可靠性评价等级结果', hzb_8, zzb_8, len(data_8), len(data_8[0]),
                                    data_8)

            doc.add_paragraph().add_run(f'根据计算得到的综合确定度，按照隶属度最大原则得到最终的隶属度等级为{level}')

            options = QFileDialog.Options()
            file_name, _ = QFileDialog.getSaveFileName(caption='Save File', directory='',
                                                       filter="Word Files (*.docx);;All Files (*)", options=options)

            doc.save(file_name)
            main_win.lable_sj_lsave.setText("导出成功！")

        except:
            main_win.lable_sj_lsave.setText("导出失败，请按照要求完成操作！")
        return


    main_win.sj_w.clicked.connect(
        lambda: {sj_w.show()}
    )
    main_win.qsm_w.clicked.connect(
        lambda: {qsm_w.show()}
    )

    # 权重计算
    # 设计阶段
    sj_w.pushButton_s1.clicked.connect(
        lambda: {sj_win_sub.show()}
    )

    sj_w.pushButton_s2.clicked.connect(
        lambda: {sj_win_obj.show()}
    )

    sj_w.pushButton_s3.clicked.connect(
        sj_zh_ww
    )

    sj_w.pushButton_s4.clicked.connect(
        sj_save_ww
    )
    sj_win_sub.save_button.clicked.connect(
        sj_sub_ww
    )
    sj_win_obj.save_button.clicked.connect(
        sj_obj_ww
    )
    # 全寿命阶段
    qsm_w.pushButton_s1.clicked.connect(
        lambda: {win_sub.show()}
    )

    qsm_w.pushButton_s2.clicked.connect(
        lambda: {win_obj.show()}
    )
    qsm_w.pushButton_s3.clicked.connect(
        zh_ww
    )
    qsm_w.pushButton_s4.clicked.connect(
        save_ww
    )
    win_sub.save_button.clicked.connect(
        sub_ww
    )
    win_obj.save_button.clicked.connect(
        obj_ww
    )

    # 确定度计算
    main_win.sj_n.clicked.connect(
        lambda: {sj_n.show()}
    )
    sj_n.pushButton_fig.clicked.connect(
        lambda: {table.show()}
    )
    sj_n.pushButton_save.clicked.connect(
        sj_save_nn
    )

    main_win.qsm_n.clicked.connect(
        lambda: {qsm_n.show()}
    )
    qsm_n.pushButton_fig.clicked.connect(
        lambda: {table.show()}
    )
    qsm_n.pushButton_save.clicked.connect(
        save_nn
    )

    # 可靠性评价
    main_win.qsm_level.clicked.connect(
        res
    )
    main_win.qsm_save.clicked.connect(
        save
    )
    main_win.sj_level.clicked.connect(
        sj_res
    )
    main_win.sj_save.clicked.connect(
        sj_save
    )
    sys.exit(app.exec_())
